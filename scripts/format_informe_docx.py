#!/usr/bin/env python3
"""Reformatea el informe MahjongIA como trabajo de investigación (Word)."""

from __future__ import annotations

import re
import shutil
import zipfile
from datetime import date
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SOURCE = Path("/home/ricar/Downloads/MahjongIA - .docx")
OUTPUT = Path("/home/ricar/Downloads/MahjongIA - Informe Final (formato investigación).docx")
BACKUP = Path("/home/ricar/Downloads/MahjongIA - original backup.docx")

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
HEADING1_SIZE = Pt(14)
HEADING2_SIZE = Pt(12)
CODE_FONT = "Courier New"


def extract_paragraphs(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return [
        "".join(t.text or "" for t in p.iter(f"{w}t"))
        for p in root.iter(f"{w}p")
    ]


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def set_run_font(run, *, size=BODY_SIZE, bold=False, italic=False, name=FONT) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=0.0) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(indent_cm) if indent_cm else None


def add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)
        set_run_font(run, size=Pt(10))


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    note = doc.add_paragraph()
    nr = note.add_run(
        "(En Microsoft Word: clic derecho sobre el índice → Actualizar campo → Actualizar toda la tabla.)"
    )
    set_run_font(nr, size=Pt(10), italic=True)
    set_paragraph_format(note, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_body(doc: Document, text: str, *, indent=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r)
    set_paragraph_format(p, indent_cm=1.25 if indent else 0.0)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=HEADING1_SIZE if level == 1 else HEADING2_SIZE,
            bold=True,
        )
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, size=Pt(10), name=CODE_FONT)
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.left_indent = Cm(1.0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=Pt(11))
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=Pt(11))
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    def center_line(text: str, *, size=BODY_SIZE, bold=False) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_after = Pt(10)

    for _ in range(3):
        doc.add_paragraph()
    center_line("[Nombre de la institución educativa]", size=Pt(12))
    center_line("Inteligencia Artificial II", size=Pt(12))
    doc.add_paragraph()
    center_line(
        "Agente de inteligencia artificial para Riichi Mahjong\n"
        "mediante aprendizaje por imitación y despliegue en RiichiLabs",
        size=Pt(16),
        bold=True,
    )
    doc.add_paragraph()
    center_line("Informe de investigación", size=Pt(14), bold=True)
    doc.add_paragraph()
    center_line("Autor(es): [Nombre del estudiante]", size=Pt(12))
    center_line("Asesor(a): [Nombre del asesor]", size=Pt(12))
    center_line(f"Fecha: {date.today().strftime('%d de %B de %Y')}", size=Pt(12))
    doc.add_page_break()


def fix_abstract(text: str) -> str:
    text = text.replace("alcanza X%", "alcanza aproximadamente 69,16 % de precisión top-1")
    text = text.replace("Introduccion", "Introducción")
    return text


def build_document(raw_paras: list[str]) -> Document:
    doc = Document()
    set_margins(doc)

    # Estilos base del documento
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    add_cover(doc)

    add_heading(doc, "Índice general", 1)
    add_toc_field(doc)
    doc.add_page_break()

    # --- Resumen ---
    add_heading(doc, "Resumen", 1)
    abstract = fix_abstract(
        next(
            (
                p
                for p in raw_paras
                if p.strip()
                and not p.startswith("Mahjong")
                and "desarrolló un agente" in p
            ),
            raw_paras[3] if len(raw_paras) > 3 else "",
        )
    )
    add_body(doc, abstract, indent=True)
    kw = doc.add_paragraph()
    kr = kw.add_run(
        "Palabras clave: Riichi Mahjong, aprendizaje por imitación, MJAI, "
        "redes neuronales, RiichiEnv, RiichiLabs."
    )
    set_run_font(kr, italic=True)
    set_paragraph_format(kw, align=WD_ALIGN_PARAGRAPH.LEFT)

    # --- 1 Introducción ---
    add_heading(doc, "1. Introducción", 1)
    intro_parts = []
    capture = False
    for p in raw_paras:
        if p.strip() == "Introduccion":
            capture = True
            continue
        if capture and p.strip() == "Planteamiento del problema y objetivos":
            break
        if capture and p.strip():
            intro_parts.append(p.strip())
    for para in intro_parts:
        add_body(doc, para, indent=True)

    # --- 2 Planteamiento ---
    add_heading(doc, "2. Planteamiento del problema y objetivos", 1)
    add_heading(doc, "2.1 Planteamiento del problema", 2)
    problem = []
    capture = False
    for p in raw_paras:
        if p.strip() == "Planteamiento del problema y objetivos":
            capture = True
            continue
        if capture and p.strip() == "Objetivo general":
            break
        if capture and p.strip():
            if p.startswith(("Espacio de acciones", "Desbalance", "Evaluación multidimensional", "Robustez")):
                add_bullet(doc, p)
            else:
                problem.append(p)
    for para in problem:
        add_body(doc, para, indent=True)

    add_heading(doc, "2.2 Objetivo general", 2)
    for p in raw_paras:
        if "Desarrollar, entrenar y evaluar un agente" in p:
            add_body(doc, p.strip(), indent=True)
            break

    add_heading(doc, "2.3 Objetivos específicos", 2)
    objs = [
        "Implementar un pipeline de procesamiento de datos desde replays MJAI.",
        "Diseñar una representación numérica del estado del juego.",
        "Entrenar modelos de política supervisados (baseline, MLP y convolucional).",
        "Incorporar mejoras orientadas a descartes y calidad del aprendizaje.",
        "Evaluar el desempeño del agente en modos offline, local y ranked.",
    ]
    for o in objs:
        add_bullet(doc, o)

    # --- 3 Dataset ---
    add_heading(doc, "3. Descripción del conjunto de datos", 1)
    add_body(
        doc,
        "El conjunto de datos utilizado es un corpus de partidas en protocolo MJAI. "
        "Cada archivo almacena la secuencia de eventos desde el reparto inicial hasta "
        "las puntuaciones finales. Las repeticiones provienen de partidas de alto nivel "
        "en Tenhou.net, plataforma de referencia en Japón.",
        indent=True,
    )
    add_heading(doc, "3.1 Formatos y estructura", 2)
    add_bullet(doc, "Formatos de texto JSON: .json, .jsonl")
    add_bullet(doc, "Variantes de logs MJAI: .mjson, .mjai")
    add_bullet(doc, "Formatos comprimidos: .gz, .zip")
    add_body(
        doc,
        "Cada línea describe un evento con campo type (start_game, tsumo, dahai, pon, etc.) "
        "y metadatos como actor o ficha (pai).",
        indent=True,
    )
    add_code_block(
        doc,
        [
            '{"type":"start_game","names":["P1","P2","P3","P4"],"kyoku_first":0,"aka_flag":true}',
            '{"type":"tsumo","actor":0,"pai":"S"}',
            '{"type":"dahai","actor":0,"pai":"S","tsumogiri":true}',
            '{"type":"pon","actor":1,"target":2,"pai":"S","consumed":["S","S"]}',
        ],
    )

    # --- 4 Metodología ---
    add_heading(doc, "4. Metodología", 1)
    add_heading(doc, "4.1 Preprocesamiento de datos", 2)
    add_body(
        doc,
        "El procesamiento inicia con el parseo de replays y la reconstrucción del juego "
        "mediante RiichiEnv. Solo se conservan decisiones con acciones legales consistentes "
        "con el motor de reglas. Las observaciones se codifican como tensores float32 de forma "
        "(74, 34) en modo base o (215, 34) extendido. Se construye un ActionVocabulary dinámico "
        "y una máscara de acciones legales por turno.",
        indent=True,
    )
    add_body(
        doc,
        "La partición entrenamiento/validación se realiza con split_replay_paths "
        "(ratio por defecto 90 % / 10 %) o con hold-out temporal en data/2025.",
        indent=True,
    )

    add_heading(doc, "4.2 Selección del modelo", 2)
    add_bullet(
        doc,
        "Baseline (action-prior): frecuencias marginales sin red neuronal; referencia inferior.",
    )
    add_bullet(
        doc,
        "MLP: tensor aplanado → capas densas; rápido pero sin explotar estructura 34×canales.",
    )
    add_bullet(
        doc,
        "Conv1D: dos capas convolucionales sobre canales; arquitectura adoptada en v4/v5.",
    )
    add_bullet(
        doc,
        "policy_with_discard_head (v5): política global + especialista en 34 tipos de descarte.",
    )

    add_heading(doc, "4.3 Entrenamiento del modelo", 2)
    add_body(doc, "Comando principal de entrenamiento del modelo v5:", indent=True)
    add_code_block(
        doc,
        [
            "mahjong-ai-train \\",
            "  --train-data data \\",
            "  --validation-data data/2025 \\",
            "  --output models/allyears_v5.pt \\",
            "  --epochs 5 \\",
            "  --extended \\",
            "  --model-arch conv \\",
            "  --model-type policy-network \\",
            "  --action-type-weight-power 0 \\",
            "  --riichi-weight-multiplier 6.0 \\",
            "  --train-discard-head \\",
            "  --example-weighting \\",
            "  --skip-bad-replays \\",
            "  --max-examples 250000",
        ],
    )
    add_bullet(doc, "Escaneo inicial: construcción del ActionVocabulary.")
    add_bullet(doc, "Iteración por batches con pérdida ponderada (cross-entropy enmascarada).")
    add_bullet(doc, "Validación periódica en data/2025 con early stopping.")
    add_bullet(doc, "Persistencia del mejor checkpoint y archivo .metrics.json.")

    # --- 5 Evaluación ---
    add_heading(doc, "5. Evaluación del desempeño", 1)
    add_body(
        doc,
        "La evaluación se organiza en capas complementarias: imitación offline, "
        "simulador local (RiichiEnv) y partidas en línea (RiichiLabs). "
        "Una alta precisión de imitación no implica por sí sola un buen ranking en mesa.",
        indent=True,
    )

    add_heading(doc, "5.1 Evaluación offline", 2)
    add_code_block(
        doc,
        [
            "mahjong-ai-evaluate offline \\",
            "  --data data/2025 \\",
            "  --model models/allyears_v5.pt \\",
            "  --max-examples 100000 \\",
            "  > results/allyears_v5_offline.json",
        ],
    )
    add_table(
        doc,
        ["Métrica", "Valor v5", "Interpretación"],
        [
            ["Top-1 global", "69,16 %", "Coincidencia estricta con el experto"],
            ["Top-5 global", "96,15 %", "Jugada humana en el top-5 legal"],
            ["Acciones fuera de vocabulario", "0,011 %", "Desajuste mínimo simulador/checkpoint"],
            ["raw_top1_illegal_rate", "~99,95 %", "Enmascaramiento legal indispensable"],
        ],
    )
    add_table(
        doc,
        ["Tipo", "Ejemplos", "Top-1", "Top-5"],
        [
            ["Global", "99 999", "69,16 %", "96,15 %"],
            ["DISCARD", "93 030", "67,21 %", "95,89 %"],
            ["PON", "2 175", "99,59 %", "100 %"],
            ["CHI", "1 573", "89,45 %", "99,87 %"],
            ["RIICHI", "1 418", "95,42 %", "100 %"],
            ["ANKAN", "86", "39,53 %", "75,58 %"],
        ],
    )
    add_table(
        doc,
        ["Indicador de descarte", "Modelo v5", "Experto"],
        [
            ["discard_match_rate", "68,34 %", "—"],
            ["model_shanten_regression_rate", "5,54 %", "0,07 %"],
            ["model_red_five_discard_rate", "0,53 %", "0,69 %"],
        ],
    )

    add_heading(doc, "5.2 Evaluación local", 2)
    add_code_block(
        doc,
        [
            "mahjong-ai-evaluate local \\",
            "  --model models/allyears_v5.pt \\",
            "  --games 100 \\",
            "  --opponent random \\",
            "  --game-mode 4p-red-single",
        ],
    )
    add_table(
        doc,
        ["Experimento", "Partidas", "Oponente", "mean_rank", "1.er puesto", "fallback_rate"],
        [
            ["v5 vs random (asiento 0)", "100", "random", "1,01", "99 %", "0 %"],
            ["v5 vs random (rotación)", "100", "random", "1,62", "46 %", "0 %"],
            ["v5 vs fallback", "100", "fallback", "1,68", "65 %", "0 %"],
            ["Solo fallback", "100", "random", "1,02", "98 %", "100 %"],
        ],
    )
    add_body(
        doc,
        "Interpretación cautelosa: frente a oponente random los rankings pueden ser optimistas; "
        "la rotación de asiento (--all-seats) es esencial para comparar agentes con equidad. "
        "Frente a oponente fallback, v5 mantiene ventaja moderada (65 % de primeros puestos).",
        indent=True,
    )

    add_heading(doc, "5.3 Evaluación en línea (RiichiLabs)", 2)
    add_code_block(
        doc,
        [
            "set -a && source .env && set +a",
            "python scripts/run_ranked_games.py \\",
            "  --games 10 \\",
            "  --endpoint wss://game.riichi.dev/ws/ranked \\",
            "  --model models/allyears_v5.pt \\",
            "  --output results/ranked_runs.jsonl",
        ],
    )
    add_body(
        doc,
        "Las primeras partidas de prueba utilizaron versiones anteriores del modelo. "
        "Con la versión final (filtros de descarte, pesos RIICHI y example weighting), "
        "el bot fue competitivo en la plataforma, aunque en ranked se ubicó en torno a "
        "1300 puntos ELO debido a la alta competencia y baja población de jugadores.",
        indent=True,
    )

    # --- 6 Implementación ---
    add_heading(doc, "6. Implementación y aplicación del modelo", 1)
    add_table(
        doc,
        ["Comando", "Módulo", "Función"],
        [
            ["mahjong-ai-train", "training/train.py", "Entrenamiento supervisado"],
            ["mahjong-ai-evaluate", "evaluation/evaluate.py", "Evaluación offline/local"],
            ["mahjong-ai-bot", "bot_runner.py", "Bot WebSocket RiichiLabs"],
        ],
    )
    add_body(
        doc,
        "El checkpoint incluye vocabulario de acciones, esquema de features y pesos de red. "
        "En inferencia se aplica enmascaramiento legal, enrutamiento a cabeza de descarte (v5), "
        "filtros de shanten/red five y capa SafeAgent con FallbackAgent ante errores.",
        indent=True,
    )
    add_code_block(
        doc,
        [
            "mahjong-ai-bot \\",
            "  --model models/allyears_v5.pt \\",
            "  --device auto \\",
            "  --endpoint wss://game.riichi.dev/ws/validate",
        ],
    )

    # --- 7 Conclusiones ---
    add_heading(doc, "7. Análisis de resultados y conclusiones", 1)
    add_heading(doc, "7.1 Resultados principales", 2)
    add_bullet(doc, "Pipeline MJAI → modelo → evaluación → bot operativo, reproducible y modular.")
    add_bullet(doc, "Imitación offline sólida: 69,16 % top-1 y 96,15 % top-5 en hold-out 2025.")
    add_bullet(doc, "Mejora clara v3 → v5 (+7,7 pp top-1); v4 y v5 equivalentes en agregado global.")
    add_bullet(doc, "Despliegue estable: fallback_rate = 0 % en evaluación local documentada.")
    add_bullet(
        doc,
        "Brecha imitación–competitividad: ranking local depende del oponente; ranked en línea ~1300 ELO.",
    )

    add_heading(doc, "7.2 Limitaciones", 2)
    add_bullet(doc, "Sesgo de imitación y desbalanceo hacia descartes (~93 % de ejemplos).")
    add_bullet(doc, "Métricas inestables en acciones raras (ANKAN, KAKAN).")
    add_bullet(doc, "Evaluación local con random no representa fuerza ante bots fuertes (Mortal, Akochan).")

    add_heading(doc, "7.3 Trabajo futuro", 2)
    add_bullet(doc, "Ampliar corpus y balancear clases de acción.")
    add_bullet(doc, "Fine-tuning con REINFORCE contra oponentes fuertes.")
    add_bullet(doc, "Torneos sistemáticos y más partidas ranked para reducir varianza.")

  # --- Referencias ---
    add_heading(doc, "Referencias bibliográficas", 1)
    refs = [
        p.strip()
        for p in raw_paras
        if p.strip().startswith(
            ("Bijay", "critter", "Equim", "GeeksforGeeks", "Jouzu", "Nikke", "Nitasurin", "Ozaki", "Tenhou", "tsubakisakura")
        )
        or ("GitHub" in p and "Accessed" in p)
        or p.strip().startswith("---.")
    ]
    if not refs:
        refs = [p for p in raw_paras if "Accessed" in p or "github.com" in p.lower()]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        set_run_font(r, size=Pt(11))
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.left_indent = Cm(1.25)
        pf.first_line_indent = Cm(-1.25)
        pf.space_after = Pt(4)

    add_page_number_footer(doc)
    return doc


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"No se encontró el archivo: {SOURCE}")

    if not BACKUP.exists():
        shutil.copy2(SOURCE, BACKUP)

    raw = extract_paragraphs(SOURCE)
    doc = build_document(raw)
    doc.save(OUTPUT)
    print(f"Guardado: {OUTPUT}")
    print(f"Respaldo del original: {BACKUP}")


if __name__ == "__main__":
    main()
